"""Extract raw text from uploaded resume files (PDF / DOCX / TXT)."""
from pathlib import Path

import docx
import pdfplumber


def extract_text(file_path: str | Path) -> str:
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(f"Unsupported resume file type: {suffix}")


def extract_text_from_bytes(data: bytes, filename: str) -> str:
    """Used by the Streamlit uploader, which hands us bytes rather than a path."""
    import io

    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    if suffix == ".docx":
        document = docx.Document(io.BytesIO(data))
        return "\n".join(p.text for p in document.paragraphs)
    if suffix == ".txt":
        return data.decode("utf-8", errors="ignore")
    raise ValueError(f"Unsupported resume file type: {suffix}")


def _extract_pdf(path: Path) -> str:
    with pdfplumber.open(path) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages)


def _extract_docx(path: Path) -> str:
    document = docx.Document(str(path))
    return "\n".join(p.text for p in document.paragraphs)
